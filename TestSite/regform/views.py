from django.shortcuts import render
from django.http import HttpResponse
from django.core.mail import EmailMessage
from .forms import RegForm
import io
import docx
from docx.shared import Inches
from django.shortcuts import redirect




# Create your views here.
def regform(request):
	form = RegForm()
	return render(
		request,
		'regform/regform.html',
		{'form': form}
	)

def generate(title, content):
	doc = docx.Document()
	docx_title=title
	doc.add_paragraph("Фамилия: {}".format(content['surname']))
	doc.add_paragraph("Имя: {}".format(content['name']))
	doc.add_paragraph("Отчество: {}".format(content['midname']))
	doc.add_paragraph("Дата рождения: {}".format(content['birthday']))
	doc.add_paragraph("Образование: {}".format(content['obraz']))
	doc.add_paragraph("Общежитие: {}".format(content['hostel']))

	doc.add_heading("Паспорт")
	doc.add_paragraph("Серия: {}".format(content['ps']))
	doc.add_paragraph("Номер: {}".format(content['pn']))
	doc.add_paragraph("Код подразделения: {}".format(content['pc']))
	doc.add_paragraph("Кем выдан: {}".format(content['pi']))
	doc.add_paragraph("Когда выдан: {}".format(content['pd']))
	doc.add_picture(content['pg'], width=Inches(6.0))
	doc.add_picture(content['pp'], width=Inches(6.0))

	doc.add_heading("Аттестат")
	doc.add_paragraph("Номер: {}".format(content['an']))
	doc.add_paragraph("Дата выдачи: {}".format(content['ad']))
	doc.add_paragraph("Кем выдан: {}".format(content['ai']))
	doc.add_picture(content['ag'], width=Inches(6.0))
	doc.add_picture(content['am'], width=Inches(6.0))

	doc.add_paragraph("Телефон для связи: {}".format(content['phone']))
	
	doc.add_picture(content['statement'], width=Inches(6.0))
	doc.add_picture(content['agreement'], width=Inches(6.0))
	# doc.save('doc.docx')
	f = io.BytesIO()
	doc.save(f)
	return f

def reg(request):
	if request.method == 'POST':
		form = RegForm(request.POST, request.FILES)
		# print('Form:\n', form)
		# print("Post:\n", request.POST, "\n===================================")
		# print("Files:\n", request.FILES, "\n===================================")
		# print(form.is_valid())
		if form.is_valid():
			cd = form.cleaned_data

			print('Cleaned_data:\n', cd)
			docattachment = generate('Test', cd)
			msg = EmailMessage(
				cd['surname'],
				cd['surname'] +' ' + cd['name'] + ' ' + cd['midname'],
				'priem@aviacollege.ru',
				['rak.priem@yandex.ru']
			)
			
			msg.attach("Test.doc",docattachment.getvalue(),'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
			msg.send()
			
			# return HttpResponse("Анкета успешно отправлена")
			response = redirect('/success/')
			return response
		return HttpResponse("common")
	else:
		return HttpResponse("common")

def success(request):
	return render(
		request,
		'regform/message.html',
		{'msg': 'Анкета успешно отправлена'}
	)